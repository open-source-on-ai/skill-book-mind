#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""book_to_mindmap.py — 多格式书籍批量转思维导图大纲（mindmap 格式 Markdown 大纲）。

支持格式：
    .pdf   PDF          —— 按 --pages 页码范围选取（1 起始）
    .docx  Word         —— 按 Word 标题样式或「第X章/Chapter N」识别章节，--chapters 选取
    .epub  EPUB         —— 按 spine 阅读顺序切章，--chapters 选取
    .txt / .md 纯文本   —— 按「第X章/Chapter N」识别章节，--chapters 选取
    .mobi / .azw3       —— 尽力支持：需 pip install mobi；未安装时提示改用 Calibre 转 EPUB

流程：读取书籍 → 选取范围 → 提取文本 → 调用 GLM / Claude API
（system prompt 从外部 prompt.md 读取，不硬编码）→ 结果保存为同名 .md。
每个分块失败自动重试 3 次（指数退避 2/4/8 秒）并记录日志。

用法示例：
    python book_to_mindmap.py book.pdf  --list
    python book_to_mindmap.py book.pdf  --pages 10-25
    python book_to_mindmap.py book.epub --chapters 3-5 --chunk-pages 2
    python book_to_mindmap.py book.docx --chapters 1-2 --dry-run
    python book_to_mindmap.py book.txt

API 配置：
    GLM（默认）：  --base-url https://open.bigmodel.cn/api/paas/v4  --model glm-4.6
    Claude：      --base-url https://api.anthropic.com/v1/         --model claude-sonnet-4-5
    密钥：        --api-key 或环境变量 ZHIPUAI_API_KEY / GLM_API_KEY / OPENAI_API_KEY / ANTHROPIC_API_KEY
"""

import argparse
import logging
import os
import posixpath
import re
import shutil
import sys
import tempfile
import time
import xml.etree.ElementTree as ET
import zipfile
from html.parser import HTMLParser
from pathlib import Path

DEFAULT_BASE_URL = "https://open.bigmodel.cn/api/paas/v4"
DEFAULT_MODEL = "glm-4.6"
RETRY_TIMES = 3
CHUNK_CHAR_FALLBACK = 4000  # 无章节标题的纯文本按此字数切段

# 「第X章 / 第X节 / 第X篇 / Chapter N」章节标题识别
CHAPTER_RE = re.compile(
    r"^\s*(?:第\s*[0-9０-９一二三四五六七八九十百千万零两]+\s*[章节篇部回]|Chapter\s+\d+)",
    re.IGNORECASE,
)

log = logging.getLogger("book2mindmap")


# ---------------------------------------------------------------- 文本工具

def decode_text(data: bytes) -> str:
    """UTF-8 优先，失败回退 GBK（中文 txt 常见编码）。"""
    for enc in ("utf-8", "gb18030"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


class _HTMLTextParser(HTMLParser):
    """把 XHTML/HTML 剥成纯文本，块级标签处断行。"""

    _BREAK_TAGS = {"p", "div", "br", "li", "tr", "section", "article",
                   "h1", "h2", "h3", "h4", "h5", "h6", "blockquote"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts = []
        self._skip = 0

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style", "title"):
            self._skip += 1
        elif tag in self._BREAK_TAGS:
            self.parts.append("\n")

    def handle_endtag(self, tag):
        if tag in ("script", "style", "title"):
            self._skip = max(0, self._skip - 1)

    def handle_data(self, data):
        if not self._skip:
            self.parts.append(data)


def html_to_text(html: str) -> str:
    parser = _HTMLTextParser()
    parser.feed(html)
    text = "".join(parser.parts)
    lines = [ln.strip() for ln in text.splitlines()]
    return "\n".join(ln for ln in lines if ln)


def html_title(html: str) -> str:
    m = re.search(r"<title[^>]*>(.*?)</title>", html, re.IGNORECASE | re.DOTALL)
    if m and m.group(1).strip():
        return m.group(1).strip()[:60]
    return ""


# ---------------------------------------------------------------- 各格式加载
# 统一返回 blocks: list[dict(label=str, text=str)]，顺序即阅读顺序。

def load_pdf(path: Path):
    """v2.3 规范要求 pymupdf 优先；未安装时回退 pypdf（止损：不手写解析器）。"""
    try:
        try:
            import pymupdf as pdfmod  # pymupdf>=1.24 推荐导入名
        except ImportError:
            import fitz as pdfmod  # 旧版包名

        blocks = []
        with pdfmod.open(str(path)) as doc:
            for i, page in enumerate(doc, start=1):
                blocks.append({"label": f"第{i}页", "text": page.get_text().strip()})
        return blocks
    except ImportError:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return [{"label": f"第{i}页", "text": (page.extract_text() or "").strip()}
                for i, page in enumerate(reader.pages, start=1)]


def load_docx(path: Path):
    from docx import Document

    doc = Document(str(path))
    top_headings = {"heading 1", "heading 2", "标题 1", "标题 2"}
    blocks, buf, title = [], [], None

    def is_chapter_start(par):
        style = (par.style.name or "").strip() if par.style is not None else ""
        return style.lower() in top_headings or bool(CHAPTER_RE.match(par.text))

    for par in doc.paragraphs:
        text = par.text.strip()
        if not text:
            continue
        if is_chapter_start(par):
            if title is not None and any(ln.strip() for ln in buf):
                blocks.append({"label": title, "text": "\n".join(buf).strip()})
            title, buf = text[:60], [text]
        else:
            buf.append(text)
    if title is not None and any(ln.strip() for ln in buf):
        blocks.append({"label": title, "text": "\n".join(buf).strip()})
    if not blocks:  # 全文无标题结构：整篇作为单块
        whole = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        blocks.append({"label": path.stem, "text": whole})
    return blocks


def load_epub(path: Path):
    ns_container = "{urn:oasis:names:tc:opendocument:xmlns:container}rootfile"
    ns_opf = "{http://www.idpf.org/2007/opf}"

    with zipfile.ZipFile(str(path)) as zf:
        container = ET.fromstring(zf.read("META-INF/container.xml"))
        rootfile = container.find(f".//{ns_container}")
        if rootfile is None or not rootfile.get("full-path"):
            raise ValueError("EPUB 缺少 container.xml 中的 rootfile 声明")
        opf_path = rootfile.get("full-path")
        opf_dir = posixpath.dirname(opf_path)
        opf = ET.fromstring(zf.read(opf_path))
        manifest = {item.get("id"): item.get("href")
                    for item in opf.iter(f"{ns_opf}item") if item.get("id")}
        spine = [ref.get("idref") for ref in opf.iter(f"{ns_opf}itemref")
                 if ref.get("idref")]

        blocks = []
        for idref in spine:
            href = manifest.get(idref)
            if not href or not href.lower().endswith((".html", ".xhtml", ".htm")):
                continue
            zip_name = posixpath.normpath(posixpath.join(opf_dir, href))
            try:
                raw = zf.read(zip_name)
            except KeyError:
                continue
            html = decode_text(raw)
            text = html_to_text(html)
            if not text.strip():
                continue
            label = html_title(html) or (text.splitlines()[0][:60] if text else href)
            blocks.append({"label": label, "text": text})
    return blocks


def load_txt(path: Path):
    text = decode_text(path.read_bytes())
    lines = text.splitlines()
    blocks, buf, title = [], [], None
    for ln in lines:
        if CHAPTER_RE.match(ln):
            if title is not None and buf:
                blocks.append({"label": title, "text": "\n".join(buf).strip()})
            title, buf = ln.strip()[:60], [ln.strip()]
        else:
            if ln.strip() or (buf and buf[-1].strip()):
                buf.append(ln.rstrip())
    if title is not None and buf:
        blocks.append({"label": title, "text": "\n".join(buf).strip()})
    if not blocks:  # 无章节标题：按段落累积到约 CHUNK_CHAR_FALLBACK 字一段
        blocks, buf, size = [], [], 0
        for para in re.split(r"\n\s*\n", text):
            para = para.strip()
            if not para:
                continue
            buf.append(para)
            size += len(para)
            if size >= CHUNK_CHAR_FALLBACK:
                blocks.append({"label": f"段落块 {len(blocks) + 1}", "text": "\n".join(buf)})
                buf, size = [], 0
        if buf:
            blocks.append({"label": f"段落块 {len(blocks) + 1}", "text": "\n".join(buf)})
    return blocks


def load_mobi(path: Path):
    """MOBI/AZW3：借助 mobi 包解包（通常得到 EPUB），失败时提示转 EPUB。"""
    try:
        from mobi import extract
    except ImportError:
        raise SystemExit(
            "MOBI/AZW3 需要额外依赖：pip install mobi；"
            "或先用 Calibre 将其转换为 EPUB 后再处理。"
        )
    tmpdir, main_file = extract(str(path))
    try:
        candidates = [Path(main_file)] + [
            p for p in Path(tmpdir).rglob("*")
            if p.suffix.lower() in (".epub", ".html", ".htm")
        ]
        for cand in candidates:
            if cand.suffix.lower() == ".epub":
                return load_epub(cand)
        for cand in candidates:  # 无 EPUB 时退回 HTML 全文
            if cand.suffix.lower() in (".html", ".htm"):
                text = html_to_text(decode_text(cand.read_bytes()))
                if text.strip():
                    return [{"label": path.stem, "text": text}]
        raise SystemExit(f"无法从 {path.name} 中提取可读文本，建议用 Calibre 转换为 EPUB。")
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".epub": load_epub,
    ".txt": load_txt,
    ".md": load_txt,
    ".mobi": load_mobi,
    ".azw3": load_mobi,
}


# ---------------------------------------------------------------- 选取与分块

def parse_indices(spec: str) -> list:
    """把 '10-25' / '3,5,10-12' 展开为去重升序的 1 起始编号列表。"""
    out = []
    for part in spec.replace("，", ",").split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            lo, _, hi = part.partition("-")
            lo, hi = int(lo), int(hi)
            if lo > hi:
                lo, hi = hi, lo
            out.extend(range(lo, hi + 1))
        else:
            out.append(int(part))
    return sorted(set(out))


def build_chunks(blocks, chunk_size):
    """blocks 按 chunk_size 块一组拼接；chunk_size<=0 表示全部并为一组。"""
    if chunk_size <= 0:
        chunk_size = len(blocks) or 1
    return ["\n\n".join(b["text"] for b in blocks[i:i + chunk_size])
            for i in range(0, len(blocks), chunk_size)]


# ---------------------------------------------------------------- API 调用

def resolve_api_key(args):
    if args.api_key:
        return args.api_key
    for var in ("ZHIPUAI_API_KEY", "GLM_API_KEY", "OPENAI_API_KEY", "ANTHROPIC_API_KEY"):
        if os.environ.get(var):
            return os.environ[var]
    return ""


def call_with_retry(client, model, system_prompt, user_text):
    last_err = None
    for attempt in range(1, RETRY_TIMES + 1):
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_text},
                ],
            )
            content = (resp.choices[0].message.content or "").strip()
            if not content:
                raise ValueError("模型返回了空内容")
            return content
        except Exception as err:  # 网络/限流/空返回统一走重试
            last_err = err
            log.warning("第 %d/%d 次调用失败：%s", attempt, RETRY_TIMES, err)
            if attempt < RETRY_TIMES:
                time.sleep(2 ** attempt)
    raise RuntimeError(f"重试 {RETRY_TIMES} 次后仍失败：{last_err}")


# ---------------------------------------------------------------- 主流程

def main(argv=None):
    if hasattr(sys.stdout, "reconfigure"):  # Windows 控制台中文输出
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")

    parser = argparse.ArgumentParser(
        description="多格式书籍批量转 mindmap 格式 Markdown 大纲（PDF/DOCX/EPUB/TXT，MOBI 尽力支持）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__.split("用法示例：")[-1],
    )
    parser.add_argument("book", help="书籍文件路径（.pdf/.docx/.epub/.txt/.md/.mobi/.azw3）")
    parser.add_argument("--pages", help="PDF 专用：页码范围，如 10-25 或 3,5,10-12（1 起始）")
    parser.add_argument("--chapters", help="非 PDF 格式：章节序号范围（按 --list 显示的编号，1 起始）")
    parser.add_argument("--chunk-pages", type=int, default=0,
                        help="每 N 个块（页/章）合并为一次 API 调用；0=全部合并为一次")
    parser.add_argument("--prompt", default=str(Path(__file__).with_name("prompt.md")),
                        help="system prompt 文件路径（默认同目录 prompt.md）")
    parser.add_argument("--out", default=".", help="输出目录（默认当前目录）")
    parser.add_argument("--base-url", default=DEFAULT_BASE_URL,
                        help=f"OpenAI 兼容 API 地址（默认 GLM：{DEFAULT_BASE_URL}）")
    parser.add_argument("--model", default=DEFAULT_MODEL, help=f"模型名（默认 {DEFAULT_MODEL}）")
    parser.add_argument("--api-key", default="", help="API 密钥；缺省读环境变量")
    parser.add_argument("--list", action="store_true", help="仅列出可选取的页/章节清单后退出")
    parser.add_argument("--dry-run", action="store_true",
                        help="仅提取并预览文本，不调用 API、不写结果文件")
    args = parser.parse_args(argv)

    book = Path(args.book)
    if not book.is_file():
        parser.error(f"文件不存在：{book}")
    loader = LOADERS.get(book.suffix.lower())
    if loader is None:
        parser.error(f"不支持的格式 {book.suffix}，支持：{'/'.join(LOADERS)}")
    if args.pages and args.chapters:
        parser.error("--pages 与 --chapters 只能二选一")
    if args.pages and book.suffix.lower() != ".pdf":
        parser.error("--pages 仅用于 PDF；其他格式请用 --chapters")
    if args.chapters and book.suffix.lower() == ".pdf":
        parser.error("PDF 请用 --pages 按页码选取")

    blocks = loader(book)

    # --list：仅展示可选取单元
    if args.list:
        print(f"{book.name} 共 {len(blocks)} 个可选单元：")
        for i, b in enumerate(blocks, start=1):
            print(f"  {i:>4}  {b['label']}  ({len(b['text'])} 字)")
        return 0

    # 范围选取
    if args.pages or args.chapters:
        spec = args.pages or args.chapters
        try:
            picked = parse_indices(spec)
        except ValueError:
            parser.error(f"范围格式错误：{spec}（示例：10-25 或 3,5,10-12）")
        bad = [n for n in picked if not 1 <= n <= len(blocks)]
        if bad:
            parser.error(f"编号超出范围 {bad}，可选 1~{len(blocks)}，先用 --list 查看")
        blocks = [blocks[n - 1] for n in picked]
        unit = "p" if args.pages else "c"
        suffix = f"{unit}{min(picked)}-{max(picked)}"
    else:
        suffix = "all"

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_md = out_dir / f"{book.stem}_{suffix}.md"
    log_file = out_dir / f"{book.stem}_{suffix}.log"
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[logging.StreamHandler(sys.stderr),
                  logging.FileHandler(log_file, encoding="utf-8")],
    )
    log.info("输入：%s（%s），选取 %d 个单元，输出：%s",
             book.name, book.suffix.lower().lstrip("."), len(blocks), out_md.name)

    chunks = build_chunks(blocks, args.chunk_pages)
    total_chars = sum(len(c) for c in chunks)
    log.info("共 %d 个分块、%d 字；分块预览：%s", len(chunks), total_chars,
             chunks[0][:120].replace("\n", " ") if chunks else "(空)")

    if args.dry_run:
        print(f"[dry-run] {book.name} 选取 {len(blocks)} 个单元，"
              f"{len(chunks)} 个分块，共 {total_chars} 字。前 300 字预览：")
        print((chunks[0] if chunks else "")[:300])
        return 0

    system_prompt = Path(args.prompt).read_text(encoding="utf-8").strip()
    if not system_prompt:
        parser.error(f"system prompt 文件为空：{args.prompt}")

    api_key = resolve_api_key(args)
    if not api_key:
        parser.error("缺少 API 密钥：请设置 --api-key 或环境变量 "
                     "ZHIPUAI_API_KEY / GLM_API_KEY / OPENAI_API_KEY / ANTHROPIC_API_KEY")
    try:
        from openai import OpenAI
    except ImportError:
        parser.error("缺少依赖：pip install openai")

    client = OpenAI(base_url=args.base_url, api_key=api_key)
    log.info("API：%s，模型：%s，system prompt：%s（%d 字）",
             args.base_url, args.model, args.prompt, len(system_prompt))

    results = []
    for idx, chunk in enumerate(chunks, start=1):
        log.info("处理分块 %d/%d（%d 字）…", idx, len(chunks), len(chunk))
        try:
            results.append(call_with_retry(client, args.model, system_prompt, chunk))
        except RuntimeError as err:
            log.error("分块 %d 失败：%s", idx, err)
            if results:
                out_md.write_text("\n".join(results) + "\n", encoding="utf-8")
                log.error("已保存此前完成的 %d 个分块到 %s", len(results), out_md.name)
            return 1

    # 规范要求：多章连续输出，章间不加分隔语
    out_md.write_text("\n".join(r.strip() for r in results) + "\n", encoding="utf-8")
    log.info("完成：%s（%d 字），日志：%s", out_md.name, out_md.stat().st_size, log_file.name)
    return 0


if __name__ == "__main__":
    sys.exit(main())
